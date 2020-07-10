#Prpgram to create a new Word document and save the inputs in it given by the user.


#This is a third party module install it by typing 'pip install python-docx'

#import the 'python-docx'
import docx

#inputs from the user
print('Enter your name: ')
name = input()

print('Enter your age: ')
age = input()

#Opening a new document object
d = docx.Document()

#add the respective inputs in new paragraphs
d.add_paragraph(name)
d.add_paragraph(age)

#save the file
d.save('file_Location/file_name.docx')
